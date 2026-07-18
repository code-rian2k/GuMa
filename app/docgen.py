"""
Erstellt Word-Dokumente (Anschreiben, Gutachten-Grundgerüst) aus einer von der
Nutzerin selbst hinterlegten Word-Vorlage (siehe app.vorlagen), indem
{{PLATZHALTER}} durch echte Falldaten ersetzt werden. Die Formatierung der
Vorlage bleibt erhalten.
"""
import os
import re
import docx

PLATZHALTER_MUSTER = re.compile(r"\{\{[A-Z_]+\}\}")

# Diese Platzhalter füllt GuMa beim Erstellen automatisch aus - eine eigene
# Word-Vorlage muss genau diese Tokens (Groß-/Kleinschreibung beachten)
# verwenden, damit z.B. der Name an der richtigen Stelle landet.
ANSCHREIBEN_PLATZHALTER = [
    "EMPFAENGER_ANREDE", "EMPFAENGER_ANREDE_ENDUNG", "EMPFAENGER_NAME",
    "DATUM", "RICHTER_TEXT", "KINDER", "GUTACHTER_TELEFON", "GUTACHTER_NAME",
]
GUTACHTEN_PLATZHALTER = ["GERICHT", "ABTEILUNG", "DATUM", "AKTENZEICHEN"]


def _ersetze_in_paragraph(paragraph, werte: dict):
    voller_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in voller_text:
        return
    neuer_text = voller_text
    for schluessel, wert in werte.items():
        neuer_text = neuer_text.replace("{{%s}}" % schluessel, str(wert if wert is not None else ""))
    if neuer_text == voller_text:
        return
    if not paragraph.runs:
        paragraph.add_run(neuer_text)
        return
    paragraph.runs[0].text = neuer_text
    for run in paragraph.runs[1:]:
        run.text = ""


def fuelle_dokument(vorlage_pfad: str, ausgabe_pfad: str, werte: dict):
    dokument = docx.Document(vorlage_pfad)
    for paragraph in dokument.paragraphs:
        _ersetze_in_paragraph(paragraph, werte)
    for table in dokument.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _ersetze_in_paragraph(paragraph, werte)
    os.makedirs(os.path.dirname(ausgabe_pfad), exist_ok=True)
    dokument.save(ausgabe_pfad)
    return ausgabe_pfad


def offene_platzhalter(pfad: str):
    """Findet {{...}}-Platzhalter, die noch nicht ersetzt wurden - zur Kontrolle."""
    dokument = docx.Document(pfad)
    gefunden = set()
    for paragraph in dokument.paragraphs:
        for treffer in PLATZHALTER_MUSTER.findall(paragraph.text):
            gefunden.add(treffer)
    return gefunden


def anschreiben_erstellen(fall: dict, einstellungen: dict, ausgabe_pfad: str, vorlage_pfad: str, richter_text: str = ""):
    anrede = fall.get("empfaenger_anrede", "Frau")
    endung = "r" if anrede == "Herr" else ""
    werte = {
        "EMPFAENGER_ANREDE": anrede,
        "EMPFAENGER_ANREDE_ENDUNG": endung,
        "EMPFAENGER_NAME": fall.get("empfaenger_name", ""),
        "DATUM": fall.get("datum", ""),
        "RICHTER_TEXT": richter_text or fall.get("richter", ""),
        "KINDER": fall.get("kinder", ""),
        "GUTACHTER_TELEFON": einstellungen.get("telefon", ""),
        "GUTACHTER_NAME": einstellungen.get("name", ""),
    }
    return fuelle_dokument(vorlage_pfad, ausgabe_pfad, werte)


def gutachten_erstellen(fall: dict, ausgabe_pfad: str, vorlage_pfad: str):
    werte = {
        "GERICHT": fall.get("gericht", ""),
        "ABTEILUNG": fall.get("abteilung", ""),
        "DATUM": fall.get("datum", ""),
        "AKTENZEICHEN": fall.get("aktenzeichen", ""),
    }
    return fuelle_dokument(vorlage_pfad, ausgabe_pfad, werte)
