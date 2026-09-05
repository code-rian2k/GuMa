"""
Funktionaler 'Durchklick-Test' der GUI-Logik ohne echtes Display: simuliert
tkinter minimal (siehe tk_stub.py), führt aber die ECHTEN Methoden von
app.gui.Anwendung aus. Deckt genau die Art von Fehlern ab, die beim
manuellen Testen durch die Nutzerin aufgetreten sind (z.B. "Neuer Fall"
verschwindet bei aktivem Suchfilter).
"""
import os
import sys
import shutil
import tempfile
import unittest
import docx

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from tests.tk_stub import build_stub_tkinter
build_stub_tkinter()

import app.db as db_modul
from app import docgen as docgen_modul
from app import vorlagen as vorlagen_modul


def _test_vorlage_bauen(pfad, platzhalter):
    dokument = docx.Document()
    for name in platzhalter:
        dokument.add_paragraph("{{%s}}" % name)
    dokument.save(pfad)


class TestGuiDurchklick(unittest.TestCase):
    def setUp(self):
        self.tmp_dir = tempfile.mkdtemp()
        db_modul.DB_PATH = os.path.join(self.tmp_dir, "test.db")

        import importlib
        import app.repo as repo_modul
        import app.gui as gui_modul
        importlib.reload(repo_modul)
        importlib.reload(gui_modul)
        self.gui_modul = gui_modul

        # Speicherort für Fälle/Dokumente auf einen Testordner umleiten, BEVOR
        # die Anwendung gestartet wird - Anwendung.__init__ liest den Ordner
        # aus den Einstellungen (Standardverhalten seit dem konfigurierbaren
        # Speicherort).
        db_modul.init_db()
        repo_modul.einstellung_setzen("dokumente_ordner", os.path.join(self.tmp_dir, "dokumente"))

        # Auch der Vorlagen-Ordner (app.vorlagen) hängt an BASIS_ORDNER - für
        # Tests auf einen Testordner umleiten, damit keine Dateien im echten
        # Projektordner landen.
        gui_modul.BASIS_ORDNER = self.tmp_dir

        self.app = gui_modul.Anwendung()

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _vorlage_hinzufuegen(self, typ):
        """Legt eine Testvorlage des angegebenen Typs an und macht sie im
        Dokumente-Tab auswählbar - entspricht dem, was die Nutzerin über
        Einstellungen → Vorlagen tut."""
        platzhalter = {
            "anschreiben": docgen_modul.ANSCHREIBEN_PLATZHALTER,
            "gutachten": docgen_modul.GUTACHTEN_PLATZHALTER,
        }[typ]
        quelle = os.path.join(self.tmp_dir, f"{typ}_quelle.docx")
        _test_vorlage_bauen(quelle, platzhalter)
        vorlagen_modul.vorlage_hinzufuegen(self.gui_modul.BASIS_ORDNER, typ, f"Test-{typ}", quelle)
        self.app._vorlagen_dropdown_aktualisieren()

    def test_neuer_fall_ohne_suchtext_wird_ausgewaehlt(self):
        self.app._neuer_fall()
        self.assertIsNotNone(self.app.aktueller_fall_id)
        self.assertIn(str(self.app.aktueller_fall_id), self.app.fall_baum.get_children())

    def test_neuer_fall_MIT_aktivem_suchtext_funktioniert_trotzdem(self):
        """
        Genau der real gemeldete Fehler: Nutzerin tippt etwas ins Suchfeld
        und klickt dann 'Neuer Fall'. Der neue Fall MUSS trotzdem erscheinen
        und auswählbar sein.
        """
        self.app.suche_var.set("irgendein Suchtext, der zu nichts passt")
        self.app._neuer_fall()
        self.assertIsNotNone(self.app.aktueller_fall_id, "aktueller_fall_id darf nach Neuer Fall nicht None sein")
        self.assertIn(str(self.app.aktueller_fall_id), self.app.fall_baum.get_children())
        # Suchfeld sollte geleert worden sein, damit der Fall sichtbar ist
        self.assertEqual(self.app.suche_var.get(), "")

    def test_falldaten_speichern_nach_neuem_fall(self):
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("1 F 123/26")
        self.app.stamm_vars["in_sachen"].set("Mustermann ./. Musterfrau")
        self.app._stammdaten_speichern()

        from app import repo
        fall = repo.fall_holen(self.app.aktueller_fall_id)
        self.assertEqual(fall["aktenzeichen"], "1 F 123/26")
        self.assertEqual(fall["in_sachen"], "Mustermann ./. Musterfrau")

    def test_speichern_ohne_ausgewaehlten_fall_zeigt_warnung_statt_nichts_zu_tun(self):
        self.app.aktueller_fall_id = None
        self.app._stammdaten_speichern()
        from tkinter import messagebox
        self.assertTrue(len(messagebox._calls["showwarning"]) >= 1)

    def test_fallwechsel_mit_ungespeicherten_aenderungen_fragt_nach(self):
        from tkinter import messagebox

        self.app._neuer_fall()
        erster_fall_id = self.app.aktueller_fall_id
        self.app.stamm_vars["aktenzeichen"].set("1 F 1/26")
        self.app._stammdaten_speichern()

        self.app._neuer_fall()
        zweiter_fall_id = self.app.aktueller_fall_id
        self.assertNotEqual(erster_fall_id, zweiter_fall_id)

        # Ungespeicherte Änderung am zweiten Fall vornehmen
        self.app.stamm_vars["aktenzeichen"].set("2 F 2/26 (nicht gespeichert)")
        self.assertTrue(self.app._stammdaten_dirty)

        # "Nein" auf die Rückfrage -> Wechsel wird verhindert, Auswahl bleibt beim zweiten Fall
        messagebox._antwort_ja_nein["value"] = False
        self.app.fall_baum.selection_set(str(erster_fall_id))
        self.app._fall_ausgewaehlt()
        self.assertEqual(self.app.aktueller_fall_id, zweiter_fall_id)
        self.assertTrue(self.app._stammdaten_dirty)

        # "Ja" auf die Rückfrage -> Wechsel klappt, Änderungen werden verworfen
        messagebox._antwort_ja_nein["value"] = True
        self.app.fall_baum.selection_set(str(erster_fall_id))
        self.app._fall_ausgewaehlt()
        self.assertEqual(self.app.aktueller_fall_id, erster_fall_id)
        self.assertFalse(self.app._stammdaten_dirty)

    def test_fallwechsel_ohne_aenderungen_fragt_nicht_nach(self):
        from tkinter import messagebox

        self.app._neuer_fall()
        erster_fall_id = self.app.aktueller_fall_id
        self.app._neuer_fall()
        zweiter_fall_id = self.app.aktueller_fall_id

        anzahl_vorher = len(messagebox._calls["askyesno"])
        self.app.fall_baum.selection_set(str(erster_fall_id))
        self.app._fall_ausgewaehlt()

        self.assertEqual(self.app.aktueller_fall_id, erster_fall_id)
        self.assertEqual(len(messagebox._calls["askyesno"]), anzahl_vorher)

    def test_notiz_hinzufuegen(self):
        self.app._neuer_fall()
        self.app.neue_notiz_text.insert("1.0", "Testnotiz für den Fall")
        self.app._notiz_hinzufuegen()
        from app import repo
        notizen = repo.notizen_liste(self.app.aktueller_fall_id)
        self.assertEqual(len(notizen), 1)
        self.assertIn("Testnotiz", notizen[0]["text"])

    def test_termin_hinzufuegen(self):
        self.app._neuer_fall()
        self.app.neuer_termin_text.set("Ortstermin bei der Familie")
        self.app._termin_hinzufuegen()
        from app import repo
        termine = repo.termine_liste(self.app.aktueller_fall_id)
        self.assertEqual(len(termine), 1)

    def test_termin_an_festem_tag_fragt_nach_und_wird_bei_ablehnung_nicht_angelegt(self):
        from app import repo
        from tkinter import messagebox

        repo.einstellung_setzen("gesperrte_wochentage", "5,6")  # Sa+So
        self.app._neuer_fall()
        self.app.neuer_termin_datum.set("19.09.2026")  # ein Samstag
        self.app.neuer_termin_text.set("Ortstermin bei der Familie")

        try:
            messagebox._antwort_ja_nein["value"] = False
            self.app._termin_hinzufuegen()
        finally:
            messagebox._antwort_ja_nein["value"] = True

        self.assertGreaterEqual(len(messagebox._calls["askyesno"]), 1)
        self.assertEqual(repo.termine_liste(self.app.aktueller_fall_id), [])

    def test_termin_an_festem_tag_wird_bei_bestaetigung_angelegt(self):
        from app import repo
        from tkinter import messagebox

        repo.einstellung_setzen("gesperrte_wochentage", "5,6")  # Sa+So
        self.app._neuer_fall()
        self.app.neuer_termin_datum.set("19.09.2026")  # ein Samstag
        self.app.neuer_termin_text.set("Ortstermin bei der Familie")

        messagebox._antwort_ja_nein["value"] = True
        self.app._termin_hinzufuegen()

        self.assertEqual(len(repo.termine_liste(self.app.aktueller_fall_id)), 1)

    def test_termin_an_normalem_tag_fragt_nicht_nach(self):
        from app import repo
        from tkinter import messagebox

        repo.einstellung_setzen("gesperrte_wochentage", "5,6")  # Sa+So
        self.app._neuer_fall()
        self.app.neuer_termin_datum.set("21.09.2026")  # ein Montag
        self.app.neuer_termin_text.set("Ortstermin bei der Familie")

        anzahl_vorher = len(messagebox._calls["askyesno"])
        self.app._termin_hinzufuegen()

        self.assertEqual(len(messagebox._calls["askyesno"]), anzahl_vorher)
        self.assertEqual(len(repo.termine_liste(self.app.aktueller_fall_id)), 1)

    def test_termin_erledigt_wird_durchgestrichen_dargestellt(self):
        self.app._neuer_fall()
        self.app.neuer_termin_text.set("Ortstermin bei der Familie")
        self.app._termin_hinzufuegen()
        from app import repo
        termin_id = repo.termine_liste(self.app.aktueller_fall_id)[0]["id"]

        self.app.termine_baum.selection_set(str(termin_id))
        self.app._termin_erledigt_umschalten()
        self.assertIn("erledigt", self.app.termine_baum.item(str(termin_id))["tags"])

        # Nochmal umschalten -> wieder offen, keine Durchstreichung mehr
        self.app.termine_baum.selection_set(str(termin_id))
        self.app._termin_erledigt_umschalten()
        self.assertNotIn("erledigt", self.app.termine_baum.item(str(termin_id))["tags"])

    def test_uebersicht_fristen_ampel_ueberfaellige_frist(self):
        import datetime
        gestern = (datetime.date.today() - datetime.timedelta(days=1)).strftime("%d.%m.%Y")
        self.app._neuer_fall()
        self.app.neuer_termin_datum.set(gestern)
        self.app.neuer_termin_text.set("Überfällige Frist")
        self.app._termin_hinzufuegen()
        from app import repo
        termin_id = repo.termine_liste(self.app.aktueller_fall_id)[0]["id"]

        self.assertIn("ueberfaellig", self.app.uebersicht_fristen_baum.item(str(termin_id))["tags"])

    def test_uebersicht_fristen_ampel_bald_faellige_frist(self):
        import datetime
        morgen = (datetime.date.today() + datetime.timedelta(days=1)).strftime("%d.%m.%Y")
        self.app._neuer_fall()
        self.app.neuer_termin_datum.set(morgen)
        self.app.neuer_termin_text.set("Bald fällige Frist")
        self.app._termin_hinzufuegen()
        from app import repo
        termin_id = repo.termine_liste(self.app.aktueller_fall_id)[0]["id"]

        self.assertIn("bald_faellig", self.app.uebersicht_fristen_baum.item(str(termin_id))["tags"])

    def test_uebersicht_fristen_ampel_keine_hervorhebung_bei_weit_entfernter_frist(self):
        import datetime
        in_zwei_monaten = (datetime.date.today() + datetime.timedelta(days=60)).strftime("%d.%m.%Y")
        self.app._neuer_fall()
        self.app.neuer_termin_datum.set(in_zwei_monaten)
        self.app.neuer_termin_text.set("Frist mit viel Vorlauf")
        self.app._termin_hinzufuegen()
        from app import repo
        termin_id = repo.termine_liste(self.app.aktueller_fall_id)[0]["id"]

        self.assertEqual(self.app.uebersicht_fristen_baum.item(str(termin_id))["tags"], ())

    def test_uebersicht_fristen_zeigt_termine_ueber_alle_faelle(self):
        self.app._neuer_fall()
        self.app.neuer_termin_text.set("Ortstermin bei der Familie")
        self.app._termin_hinzufuegen()
        from app import repo
        termin_id = repo.termine_liste(self.app.aktueller_fall_id)[0]["id"]

        self.assertIn(str(termin_id), self.app.uebersicht_fristen_baum.get_children())

    def test_uebersicht_fristen_bereits_beim_start_gefuellt_ohne_fallauswahl(self):
        """Regressionstest: Beim (Neu-)Start der Anwendung ist noch kein Fall
        ausgewählt. Trotzdem müssen bereits bestehende Fristen sofort im
        Übersicht-Tab erscheinen, statt erst nach Auswahl eines Falls."""
        self.app._neuer_fall()
        self.app.neuer_termin_text.set("Bereits vorhandener Termin")
        self.app._termin_hinzufuegen()
        from app import repo
        termin_id = repo.termine_liste(self.app.aktueller_fall_id)[0]["id"]

        neue_app = self.gui_modul.Anwendung()
        self.assertIsNone(neue_app.aktueller_fall_id)
        self.assertIn(str(termin_id), neue_app.uebersicht_fristen_baum.get_children())

    def test_springe_zu_fall_waehlt_fall_aus_und_wechselt_tab(self):
        self.app._neuer_fall()
        erster_fall_id = self.app.aktueller_fall_id
        self.app._neuer_fall()

        # Aktiver Suchfilter, der den Zielfall ausblenden würde, muss geräumt werden
        self.app.suche_var.set("irgendein Text, der zu nichts passt")

        self.app._springe_zu_fall(erster_fall_id)

        self.assertEqual(self.app.aktueller_fall_id, erster_fall_id)
        self.assertEqual(self.app.suche_var.get(), "")

    def test_uebersicht_termin_oeffnen_springt_zu_fristen_tab(self):
        self.app._neuer_fall()
        fall_id = self.app.aktueller_fall_id
        self.app.neuer_termin_text.set("Ortstermin bei der Familie")
        self.app._termin_hinzufuegen()
        from app import repo
        termin_id = repo.termine_liste(fall_id)[0]["id"]

        self.app._neuer_fall()  # anderer Fall ist jetzt ausgewählt
        self.app.uebersicht_fristen_baum.selection_set(str(termin_id))
        self.app._uebersicht_termin_oeffnen()

        self.assertEqual(self.app.aktueller_fall_id, fall_id)

    def test_datei_hinzufuegen_zu_fall(self):
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("5 F 55/26")
        self.app._stammdaten_speichern()

        testdatei = os.path.join(self.tmp_dir, "foto.jpg")
        with open(testdatei, "w") as f:
            f.write("bilddaten-platzhalter")

        from tkinter import filedialog
        filedialog._werte["askopenfilenames"] = (testdatei,)
        self.app._unterlagen_hinzufuegen()

        ordner = self.app._fall_ordner()
        dateien_im_ordner = [d for d in os.listdir(ordner) if d.endswith(".jpg")]
        self.assertEqual(len(dateien_im_ordner), 1)

    def test_rechnung_anlegen_und_oeffnen_funktioniert(self):
        self.app._neuer_fall()
        self.app._rechnung_neu()
        from app import repo
        rechnungen = repo.rechnungen_fuer_fall(self.app.aktueller_fall_id)
        self.assertEqual(len(rechnungen), 1)

    def test_unterlagen_liste_aktualisiert_sich_nach_rechnungsfenster_schliessen(self):
        """Eine als Excel exportierte Rechnung landet im Fall-Ordner - die
        Unterlagen-Liste muss das ohne Fallwechsel sofort zeigen."""
        self.app._neuer_fall()
        from app import repo
        rechnung_id = repo.rechnung_anlegen(self.app.aktueller_fall_id, "01-2026", "08.07.2026")

        ordner = self.app._fall_ordner()
        self.assertEqual(self.app.unterlagen_baum.get_children(), [])

        # Simuliert die Datei, die ein Excel-Export im Rechnungsfenster erzeugt
        exportierte_datei = os.path.join(ordner, "Rechnung_01-2026.xlsx")
        with open(exportierte_datei, "w") as f:
            f.write("platzhalter")

        self.app._rechnung_fenster_oeffnen(rechnung_id)

        self.assertIn(exportierte_datei, self.app.unterlagen_baum.get_children())

    def test_rechnung_neu_nutzt_standard_saetze_aus_einstellungen(self):
        from app import repo
        repo.einstellung_setzen("standard_stundensatz", "150")
        repo.einstellung_setzen("standard_mwst_satz", "7")

        self.app._neuer_fall()
        self.app._rechnung_neu()

        rechnung = repo.rechnungen_fuer_fall(self.app.aktueller_fall_id)[0]
        self.assertEqual(rechnung["stundensatz"], 150)
        self.assertEqual(rechnung["mwst_satz"], 7)

    def test_naechste_rechnungsnummer_vermeidet_kollision_nach_loeschen(self):
        import datetime as dt
        from app import repo
        jahr = dt.date.today().year

        self.app._neuer_fall()
        self.app._rechnung_neu()  # 01
        self.app._rechnung_neu()  # 02
        self.app._rechnung_neu()  # 03
        rechnungen = repo.rechnungen_fuer_fall(self.app.aktueller_fall_id)
        self.assertEqual(len(rechnungen), 3)
        mittlere = next(r for r in rechnungen if r["rechnungsnummer"] == f"02-{jahr}")

        repo.rechnung_loeschen(mittlere["id"])

        # Ohne Fix würde hier "03-{jahr}" vorgeschlagen (Anzahl 2 + 1) und
        # mit der noch vorhandenen dritten Rechnung kollidieren.
        self.assertEqual(self.app._naechste_rechnungsnummer(), f"04-{jahr}")

    def test_fall_loeschen_entfernt_aus_liste(self):
        self.app._neuer_fall()
        fall_id = self.app.aktueller_fall_id
        self.app._fall_loeschen()
        self.assertNotIn(str(fall_id), self.app.fall_baum.get_children())

    def test_anschreiben_erstellen_ohne_vorlage_zeigt_warnung(self):
        """GuMa liefert keine Vorlagen mehr mit - ohne selbst hinzugefügte
        Vorlage darf nichts erstellt werden, sondern es muss ein klarer
        Hinweis erscheinen."""
        self.app._neuer_fall()
        self.app._anschreiben_erstellen()
        from tkinter import messagebox
        self.assertTrue(len(messagebox._calls["showwarning"]) >= 1)

    def test_anschreiben_erstellen_ueber_gui_speichert_automatisch_im_fallordner(self):
        self._vorlage_hinzufuegen("anschreiben")
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("2 F 2/26")
        self.app._stammdaten_speichern()
        self.app.empfaenger_name_var.set("Musterfrau")

        # Kein Speicherdialog mehr - es darf keine Rückfrage nötig sein
        self.app._anschreiben_erstellen()

        ordner = self.app._fall_ordner()
        erzeugte = [d for d in os.listdir(ordner) if d.startswith("Anschreiben_") and d.endswith(".docx")]
        self.assertEqual(len(erzeugte), 1)

        import docx
        text = "\n".join(p.text for p in docx.Document(os.path.join(ordner, erzeugte[0])).paragraphs)
        self.assertIn("Musterfrau", text)

    def test_anschreiben_zweimal_erstellen_ueberschreibt_nicht(self):
        self._vorlage_hinzufuegen("anschreiben")
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("8 F 8/26")
        self.app._stammdaten_speichern()
        self.app.empfaenger_name_var.set("Musterfrau")

        self.app._anschreiben_erstellen()
        self.app._anschreiben_erstellen()

        ordner = self.app._fall_ordner()
        erzeugte = [d for d in os.listdir(ordner) if d.startswith("Anschreiben_") and d.endswith(".docx")]
        self.assertEqual(len(erzeugte), 2, "Die zweite Datei darf die erste nicht überschreiben")

    def test_gutachten_erstellen_ueber_gui_speichert_automatisch_im_fallordner(self):
        self._vorlage_hinzufuegen("gutachten")
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("3 F 3/26")
        self.app.stamm_vars["gericht"].set("Augsburg")
        self.app._stammdaten_speichern()

        self.app._gutachten_erstellen()

        ordner = self.app._fall_ordner()
        pfad = os.path.join(ordner, "Gutachten.docx")
        self.assertTrue(os.path.isfile(pfad))

        import docx
        text = "\n".join(p.text for p in docx.Document(pfad).paragraphs)
        self.assertIn("3 F 3/26", text)
        self.assertIn("Augsburg", text)

    def test_uebersicht_gutachten_schnellzugriff_zeigt_erstelltes_gutachten(self):
        self._vorlage_hinzufuegen("gutachten")
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("3 F 3/26")
        self.app._stammdaten_speichern()

        self.app._gutachten_erstellen()

        eintraege = self.app.uebersicht_gutachten_baum.get_children()
        self.assertEqual(len(eintraege), 1)
        pfad = self.app._uebersicht_gutachten_pfade[eintraege[0]]
        self.assertTrue(os.path.isfile(pfad))
        self.assertTrue(os.path.basename(pfad).lower().startswith("gutachten"))

    def test_rechnungsfenster_oeffnet_und_berechnet(self):
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("4 F 4/26")
        self.app._stammdaten_speichern()

        from app import repo
        rechnung_id = repo.rechnung_anlegen(self.app.aktueller_fall_id, "01-2026", "08.07.2026")

        from app.gui_rechnung import RechnungFenster
        fall = repo.fall_holen(self.app.aktueller_fall_id)
        fenster = RechnungFenster(self.app, dict(fall), rechnung_id, self.app._fall_ordner())

        # Erste Zeitposten-Zeile mit Minuten befüllen und neu berechnen lassen
        fenster.zeitposten_zeilen[0]["minuten"].set("120")
        fenster._neu_berechnen()
        self.assertEqual(fenster._letztes_ergebnis.stunden_aufgerundet, 2)
        self.assertEqual(fenster._letztes_ergebnis.summe_zeitaufwand, 200)

        fenster._speichern()
        posten = repo.zeitposten_fuer_rechnung(rechnung_id)
        self.assertEqual(int(posten[0]["minuten"]), 120)

    def test_zusatzposten_hinzufuegen_speichern_und_berechnen(self):
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("10 F 10/26")
        self.app._stammdaten_speichern()

        from app import repo
        rechnung_id = repo.rechnung_anlegen(self.app.aktueller_fall_id, "01-2026", "08.07.2026")

        from app.gui_rechnung import RechnungFenster
        fall = repo.fall_holen(self.app.aktueller_fall_id)
        fenster = RechnungFenster(self.app, dict(fall), rechnung_id, self.app._fall_ordner())

        fenster._zusatzposten_zeile_hinzufuegen()
        fenster.zusatzposten_zeilen[0]["bezeichnung"].set("Fahrtkosten Bahn")
        fenster.zusatzposten_zeilen[0]["betrag"].set("45")
        fenster._neu_berechnen()
        self.assertEqual(fenster._letztes_ergebnis.zusatzposten_summe, 45.0)

        fenster._speichern()
        posten = repo.aufwandsposten_fuer_rechnung(rechnung_id)
        self.assertEqual(len(posten), 1)
        self.assertEqual(posten[0]["bezeichnung"], "Fahrtkosten Bahn")
        self.assertEqual(posten[0]["betrag"], 45.0)

        # Entfernen der Zeile (wie über die "x"-Schaltfläche) funktioniert ebenfalls
        eintrag = fenster.zusatzposten_zeilen[0]
        fenster.zusatzposten_zeilen.remove(eintrag)
        eintrag["frame"].destroy()
        fenster._speichern()
        self.assertEqual(repo.aufwandsposten_fuer_rechnung(rechnung_id), [])

    def test_rechnungsfenster_warnt_bei_ungespeicherten_aenderungen(self):
        from tkinter import messagebox

        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("12 F 12/26")
        self.app._stammdaten_speichern()

        from app import repo
        rechnung_id = repo.rechnung_anlegen(self.app.aktueller_fall_id, "01-2026", "08.07.2026")

        from app.gui_rechnung import RechnungFenster
        fall = repo.fall_holen(self.app.aktueller_fall_id)
        fenster = RechnungFenster(self.app, dict(fall), rechnung_id, self.app._fall_ordner())
        self.assertFalse(fenster._dirty, "nach dem Laden darf nichts als geändert markiert sein")

        fenster.stundensatz_var.set("150")
        self.assertTrue(fenster._dirty)

        # "Nein" -> Fenster bleibt offen (wird nicht zerstört)
        messagebox._antwort_ja_nein["value"] = False
        fenster._schliessen()
        self.assertTrue(fenster._dirty)

        # "Ja" -> Fenster wird geschlossen, ohne dass gespeichert wurde
        messagebox._antwort_ja_nein["value"] = True
        fenster._schliessen()
        rechnung = repo.rechnung_holen(rechnung_id)
        self.assertEqual(rechnung["stundensatz"], 100.0)  # unverändert, da nicht gespeichert

    def test_rechnungsfenster_schliesst_ohne_rueckfrage_nach_speichern(self):
        from tkinter import messagebox

        self.app._neuer_fall()
        from app import repo
        rechnung_id = repo.rechnung_anlegen(self.app.aktueller_fall_id, "01-2026", "08.07.2026")

        from app.gui_rechnung import RechnungFenster
        fall = repo.fall_holen(self.app.aktueller_fall_id)
        fenster = RechnungFenster(self.app, dict(fall), rechnung_id, self.app._fall_ordner())
        fenster.stundensatz_var.set("150")
        fenster._speichern()
        self.assertFalse(fenster._dirty)

        anzahl_vorher = len(messagebox._calls["askyesno"])
        fenster._schliessen()
        self.assertEqual(len(messagebox._calls["askyesno"]), anzahl_vorher)

    def test_rechnung_excel_export_ueber_gui(self):
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("6 F 6/26")
        self.app._stammdaten_speichern()

        from app import repo
        rechnung_id = repo.rechnung_anlegen(self.app.aktueller_fall_id, "02-2026", "08.07.2026")
        from app.gui_rechnung import RechnungFenster
        fall = repo.fall_holen(self.app.aktueller_fall_id)
        fenster = RechnungFenster(self.app, dict(fall), rechnung_id, self.app._fall_ordner())
        fenster.zeitposten_zeilen[0]["minuten"].set("60")

        ziel = os.path.join(self.tmp_dir, "Rechnung_Test.xlsx")
        from tkinter import filedialog
        filedialog._werte["asksaveasfilename"] = ziel
        fenster._exportieren()
        self.assertTrue(os.path.isfile(ziel))

    def test_backup_ueber_menu(self):
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("7 F 7/26")
        self.app._stammdaten_speichern()

        testdatei = os.path.join(self.tmp_dir, "beleg.pdf")
        with open(testdatei, "w") as f:
            f.write("pdf-platzhalter")
        from tkinter import filedialog
        filedialog._werte["askopenfilenames"] = (testdatei,)
        self.app._unterlagen_hinzufuegen()

        backup_zip = os.path.join(self.tmp_dir, "backup_ziel", "mein-backup.zip")
        filedialog._werte["asksaveasfilename"] = backup_zip
        self.app._backup_erstellen()

        self.assertTrue(os.path.isfile(backup_zip))
        import zipfile
        import app.db as db_modul
        with zipfile.ZipFile(backup_zip) as zf:
            namen = zf.namelist()
        self.assertIn(os.path.basename(db_modul.DB_PATH), namen)
        self.assertTrue(any(n.endswith("beleg.pdf") for n in namen))

    def test_backup_importieren_stellt_geloeschten_fall_wieder_her(self):
        from tkinter import filedialog
        from app import repo

        self.app._neuer_fall()
        fall_id = self.app.aktueller_fall_id
        self.app.stamm_vars["aktenzeichen"].set("8 F 8/26")
        self.app._stammdaten_speichern()

        backup_zip = os.path.join(self.tmp_dir, "backup.zip")
        filedialog._werte["asksaveasfilename"] = backup_zip
        self.app._backup_erstellen()

        repo.fall_loeschen(fall_id)
        self.assertIsNone(repo.fall_holen(fall_id))

        filedialog._werte["askopenfilename"] = backup_zip
        self.app._backup_importieren()

        wiederhergestellt = repo.fall_holen(fall_id)
        self.assertIsNotNone(wiederhergestellt)
        self.assertEqual(wiederhergestellt["aktenzeichen"], "8 F 8/26")

    def test_backup_importieren_ohne_bestaetigung_aendert_nichts(self):
        from tkinter import filedialog, messagebox
        from app import repo

        self.app._neuer_fall()
        fall_id = self.app.aktueller_fall_id
        self.app.stamm_vars["aktenzeichen"].set("9 F 9/26")
        self.app._stammdaten_speichern()

        backup_zip = os.path.join(self.tmp_dir, "backup2.zip")
        filedialog._werte["asksaveasfilename"] = backup_zip
        self.app._backup_erstellen()

        repo.fall_loeschen(fall_id)

        try:
            messagebox._antwort_ja_nein["value"] = False
            filedialog._werte["askopenfilename"] = backup_zip
            self.app._backup_importieren()
        finally:
            messagebox._antwort_ja_nein["value"] = True

        self.assertIsNone(repo.fall_holen(fall_id))  # Import wurde abgebrochen, nichts wiederhergestellt

    def test_speicherort_aendern_wird_von_der_app_uebernommen(self):
        from app import repo
        neuer_ordner = os.path.join(self.tmp_dir, "Woanders", "Fallakten")

        # Entspricht dem, was die Einstellungen-Dialog beim Speichern tut:
        repo.einstellung_setzen("dokumente_ordner", neuer_ordner)
        self.app._dokumente_ordner_neu_einlesen()

        self.assertEqual(self.gui_modul.DOKUMENTE_ORDNER, neuer_ordner)
        self.assertTrue(os.path.isdir(neuer_ordner))

        # Ein neuer Fall wird jetzt tatsächlich im neuen Ordner angelegt
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("9 F 9/26")
        self.app._stammdaten_speichern()
        fall_ordner = self.app._fall_ordner()
        self.assertTrue(fall_ordner.startswith(neuer_ordner))

    def test_einstellungen_speichern(self):
        self.app._einstellungen_oeffnen()
        from app import repo
        repo.einstellung_setzen("iban", "DE12 3456 7890 0000 0000 00")
        self.assertEqual(repo.einstellungen_holen()["iban"], "DE12 3456 7890 0000 0000 00")

    def test_einstellungen_oeffnen_mit_bereits_gesetzten_festen_tagen(self):
        from app import repo
        repo.einstellung_setzen("gesperrte_wochentage", "5,6")
        self.app._einstellungen_oeffnen()  # darf mit vorbelegten Checkboxen nicht scheitern

    def test_info_dialog_oeffnet_ohne_fehler(self):
        self.app._info_oeffnen()

    def test_suche_filtert_aber_neuer_fall_bleibt_robust(self):
        # Ersten Fall anlegen und eindeutig benennen
        self.app._neuer_fall()
        self.app.stamm_vars["aktenzeichen"].set("1 F 1/26")
        self.app.stamm_vars["in_sachen"].set("Müller")
        self.app._stammdaten_speichern()

        # Suchfeld auf einen Text setzen, der NICHT zum nächsten neuen Fall passt
        self.app.suche_var.set("Müller")
        self.assertEqual(len(self.app.fall_baum.get_children()), 1)

        # Zweiten Fall anlegen während die Suche noch aktiv "aussehen" könnte
        self.app._neuer_fall()
        self.assertIsNotNone(self.app.aktueller_fall_id)
        self.assertIn(str(self.app.aktueller_fall_id), self.app.fall_baum.get_children())

    def test_status_filter_zeigt_nur_passende_faelle(self):
        from app import repo

        fall_offen = repo.fall_anlegen({"aktenzeichen": "1 F 1/26", "status": "offen"})
        fall_bearbeitung = repo.fall_anlegen({"aktenzeichen": "2 F 2/26", "status": "in Bearbeitung"})
        self.app._faelle_neu_laden()
        self.assertEqual(len(self.app.fall_baum.get_children()), 2)

        self.app.status_filter_var.set("in Bearbeitung")

        self.assertEqual(self.app.fall_baum.get_children(), [str(fall_bearbeitung)])
        self.assertNotIn(str(fall_offen), self.app.fall_baum.get_children())

        self.app.status_filter_var.set("Alle")

        self.assertEqual(len(self.app.fall_baum.get_children()), 2)


if __name__ == "__main__":
    unittest.main()
