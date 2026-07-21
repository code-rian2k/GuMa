import os
import unittest
import tempfile
import docx

from app.docgen import (
    anschreiben_erstellen, gutachten_erstellen, offene_platzhalter,
    ANSCHREIBEN_PLATZHALTER, GUTACHTEN_PLATZHALTER,
)


def _test_vorlage_bauen(pfad, platzhalter):
    """Baut eine winzige .docx-Testvorlage, die jeden übergebenen Platzhalter
    einmal enthält - steht anstelle einer echten, von der Nutzerin selbst
    hinzugefügten Word-Vorlage (GuMa liefert keine Vorlagen mehr mit)."""
    dokument = docx.Document()
    for name in platzhalter:
        dokument.add_paragraph("{{%s}}" % name)
    dokument.save(pfad)


class TestDokumentErstellung(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.fall = {
            "empfaenger_anrede": "Frau", "empfaenger_name": "Musterfrau",
            "datum": "08.07.2026", "richter": "Herrn Müller, Richter am Amtsgericht Augsburg",
            "kinder": "Max und Lisa", "gericht": "Augsburg",
            "abteilung": "Abteilung für Familiensachen", "aktenzeichen": "123 F 456/26",
            "in_sachen": "Mustermann ./. Musterfrau", "mutter_name": "Musterfrau, Erika",
            "vater_name": "Mustermann, Max",
        }
        self.einstellungen = {
            "telefon": "0821/349 43 73", "name": "Frau Dr. Beispiel, Dipl.-Psych.",
            "absender_adresse": "Musterstraße 1, 86150 Augsburg",
        }

        self.anschreiben_vorlage = os.path.join(self.tmp, "anschreiben_vorlage.docx")
        _test_vorlage_bauen(self.anschreiben_vorlage, ANSCHREIBEN_PLATZHALTER)
        self.gutachten_vorlage = os.path.join(self.tmp, "gutachten_vorlage.docx")
        _test_vorlage_bauen(self.gutachten_vorlage, GUTACHTEN_PLATZHALTER)

    def test_anschreiben_hat_keine_offenen_platzhalter(self):
        pfad = os.path.join(self.tmp, "anschreiben.docx")
        anschreiben_erstellen(self.fall, self.einstellungen, pfad, self.anschreiben_vorlage)
        self.assertEqual(offene_platzhalter(pfad), set())

    def test_anschreiben_enthaelt_falldaten(self):
        pfad = os.path.join(self.tmp, "anschreiben.docx")
        anschreiben_erstellen(self.fall, self.einstellungen, pfad, self.anschreiben_vorlage)
        text = "\n".join(p.text for p in docx.Document(pfad).paragraphs)
        self.assertIn("Musterfrau", text)
        self.assertIn("Max und Lisa", text)
        self.assertIn("0821/349 43 73", text)
        self.assertIn("Musterstraße 1, 86150 Augsburg", text)
        self.assertIn("Mustermann ./. Musterfrau", text)
        self.assertIn("Musterfrau, Erika", text)
        self.assertIn("Mustermann, Max", text)

    def test_gutachten_hat_keine_offenen_platzhalter(self):
        pfad = os.path.join(self.tmp, "gutachten.docx")
        gutachten_erstellen(self.fall, pfad, self.gutachten_vorlage)
        self.assertEqual(offene_platzhalter(pfad), set())

    def test_gutachten_enthaelt_aktenzeichen(self):
        pfad = os.path.join(self.tmp, "gutachten.docx")
        gutachten_erstellen(self.fall, pfad, self.gutachten_vorlage)
        text = "\n".join(p.text for p in docx.Document(pfad).paragraphs)
        self.assertIn("123 F 456/26", text)
        self.assertIn("Augsburg", text)
        self.assertIn("Mustermann ./. Musterfrau", text)
        self.assertIn("Musterfrau, Erika", text)
        self.assertIn("Mustermann, Max", text)

    def test_leere_felder_erzeugen_keinen_fehler(self):
        pfad = os.path.join(self.tmp, "anschreiben_leer.docx")
        anschreiben_erstellen({}, {}, pfad, self.anschreiben_vorlage)
        self.assertTrue(os.path.exists(pfad))

    def test_unbekannter_platzhalter_wird_als_offen_gemeldet(self):
        """Enthält eine Vorlage einen Platzhalter, den GuMa nicht kennt, muss
        das erkannt werden - Grundlage für die Warnung in der GUI."""
        vorlage = os.path.join(self.tmp, "vorlage_mit_tippfehler.docx")
        _test_vorlage_bauen(vorlage, ["EMPFAENGER_NAME", "EMPFAENGER_ANEDE"])  # Tippfehler
        pfad = os.path.join(self.tmp, "anschreiben_tippfehler.docx")
        anschreiben_erstellen(self.fall, self.einstellungen, pfad, vorlage)
        self.assertEqual(offene_platzhalter(pfad), {"{{EMPFAENGER_ANEDE}}"})


if __name__ == "__main__":
    unittest.main()
